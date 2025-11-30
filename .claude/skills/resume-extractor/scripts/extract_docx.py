#!/usr/bin/env python3
"""
Extract text and structure from DOCX resumes.

Usage:
    uv run scripts/extract_docx.py <docx_path> [--output <output_file>]

Examples:
    uv run scripts/extract_docx.py resume.docx
    uv run scripts/extract_docx.py resume.docx --output extracted.txt
"""

import argparse
import sys
from pathlib import Path


def extract_docx_text(docx_path: Path) -> str:
    """Extract all text from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx not available. Run: uv sync", file=sys.stderr)
        sys.exit(1)

    doc = Document(docx_path)
    paragraphs = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def main():
    parser = argparse.ArgumentParser(description="Extract text from DOCX resume")
    parser.add_argument("docx_path", type=Path, help="Path to DOCX resume file")
    parser.add_argument("-o", "--output", type=Path, help="Output file path (default: stdout)")

    args = parser.parse_args()

    if not args.docx_path.exists():
        print(f"Error: File not found: {args.docx_path}", file=sys.stderr)
        sys.exit(1)

    if not args.docx_path.suffix.lower() == '.docx':
        print(f"Error: Not a DOCX file: {args.docx_path}", file=sys.stderr)
        sys.exit(1)

    text = extract_docx_text(args.docx_path)

    if args.output:
        args.output.write_text(text)
        print(f"Extracted text written to: {args.output}", file=sys.stderr)
    else:
        print(text)


if __name__ == "__main__":
    main()
